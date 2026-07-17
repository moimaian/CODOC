# -*- coding: utf-8 -*-
"""Final results report generator for Step 5 (CODOC Results Report, .docx).

Pulls its data from three kinds of sources, per job:
  - the .json settings snapshots written next to a job's results (docking_settings.json,
    ligand_settings.json);
  - the actions.log written into a job's CONVERSION folder as Step 2 tools are run;
  - the result folders themselves (ligand counts per CONVERSION_* subfolder, the rebuilt
    docking CSV for binding energies/RMSD/SMILES).

Kept free of PyQt so it can be unit-tested and reused without a running GUI.
"""

from __future__ import annotations

import io
import json
import os
from pathlib import Path
from typing import Any, Callable, Optional

try:
    from rdkit import Chem
    from rdkit.Chem import AllChem, Draw
    from rdkit.Chem.Draw import rdMolDraw2D
except Exception:
    Chem = None
    AllChem = None
    Draw = None
    rdMolDraw2D = None

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

COLOR_TITLE_BAR = "BDD6EE"
COLOR_SECTION_BAR = "DAE3F3"
COLOR_ACCENT_BAR = "E2EFDA"

LIGAND_ACTION_LABELS = [
    ("split_multimodel", "Split multimodel files"),
    ("split_large_folders", "Split large folders"),
    ("generate_lipinski", "Generate SMI/CSV+Lipinski"),
    ("druggability_filter", "Apply druggability filter"),
    ("move_empty", "Move empty files"),
    ("convert_pdbqt", "Convert ligands to PDBQT"),
    ("reject_pdbqt", "Reject invalid PDBQT"),
    ("recover_pdbqt", "Recover failed ligands"),
    ("fix_macrocycles", "Fix macrocycles for GPU"),
]



def require_docx() -> None:
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available in the current environment. Install python-docx to generate reports.")


def require_rdkit() -> None:
    if Chem is None or Draw is None:
        raise RuntimeError("RDKit is not available in the current environment. Install RDKit to generate reports.")


# --------------------------------------------------------------------------------------
# Filesystem-derived data
# --------------------------------------------------------------------------------------

def _count_files(path: str) -> int:
    directory = Path(path)
    if not directory.is_dir():
        return 0
    return sum(1 for entry in directory.iterdir() if entry.is_file())


def _count_pdbqt(path: str) -> int:
    directory = Path(path)
    if not directory.is_dir():
        return 0
    return sum(1 for entry in directory.glob("*.pdbqt") if entry.is_file())


def read_completed_actions(conversion_dir: str) -> set[str]:
    """Parse actions.log (written by CODOC.py each time a Step 2 tool finishes) into the
    set of operation keys that ran successfully for this job."""
    log_path = Path(conversion_dir) / "actions.log"
    if not log_path.is_file():
        return set()
    completed: set[str] = set()
    for line in log_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        parts = line.strip().split("\t")
        if len(parts) >= 3 and parts[2] == "OK":
            completed.add(parts[1])
    return completed


def compute_databank_stats(ligands_dir: str, conversion_dir: str, databank_names: list[str]) -> dict[str, dict[str, int]]:
    """Reconstruct the ligand-preparation funnel per databank straight from the folders that
    Step 2's tools actually populate - there is no single stored "initial count", so it is
    rebuilt as final + everything that was shed along the way - recoveries added back."""
    stats: dict[str, dict[str, int]] = {}
    for name in databank_names:
        final = _count_pdbqt(os.path.join(ligands_dir, name))
        rejected = _count_files(os.path.join(conversion_dir, "CONVERSION_FAILURES", name))
        empty = _count_files(os.path.join(conversion_dir, "EMPTY_LIGANDS", name))
        macrocycles_fixed = _count_files(os.path.join(conversion_dir, "MACROCYCLES", name))
        recovered = _count_files(os.path.join(ligands_dir, "RECOVERED", name))
        druggability_filtered = _count_files(os.path.join(conversion_dir, "NO_DRUGGABILITY", name))
        initial = max(final, final + rejected + empty + druggability_filtered - recovered)
        stats[name] = {
            "initial": initial,
            "rejected": rejected,
            "empty": empty,
            "macrocycles_fixed": macrocycles_fixed,
            "recovered": recovered,
            "druggability_filtered": druggability_filtered,
            "final": final,
        }
    return stats


def detect_databank_format(conversion_dir: str, ligands_dir: str, databank_name: str) -> str:
    """Best-effort original file extension for a databank (e.g. ".mol2", ".smi"), read from
    the archived originals if the conversion action ran, else from any leftover raw file."""
    for base in (
        os.path.join(conversion_dir, "CONVERSION_ORIGINALS", databank_name),
        os.path.join(ligands_dir, databank_name),
    ):
        directory = Path(base)
        if not directory.is_dir():
            continue
        counts: dict[str, int] = {}
        for entry in directory.iterdir():
            if entry.is_file() and entry.suffix.lower() != ".pdbqt":
                counts[entry.suffix.lower()] = counts.get(entry.suffix.lower(), 0) + 1
        if counts:
            return max(counts, key=lambda ext: counts[ext])
    return ""


# --------------------------------------------------------------------------------------
# 2D structure rendering
# --------------------------------------------------------------------------------------

def molecule_image_bytes(smiles: str, size: tuple[int, int] = (260, 180)) -> Optional[io.BytesIO]:
    if Chem is None or rdMolDraw2D is None or not smiles:
        return None
    try:
        mol = Chem.MolFromSmiles(smiles)
        if mol is None:
            return None
        AllChem.Compute2DCoords(mol)
        drawer = rdMolDraw2D.MolDraw2DCairo(*size)
        drawer.drawOptions().padding = 0.15
        drawer.DrawMolecule(mol)
        drawer.FinishDrawing()
        buffer = io.BytesIO(drawer.GetDrawingText())
        buffer.seek(0)
        return buffer
    except Exception:
        return None


# --------------------------------------------------------------------------------------
# docx helpers
# --------------------------------------------------------------------------------------

def _shade_cell(cell: Any, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_width(cell: Any, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def _content_width_cm(document: Any) -> float:
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 360000


def _bar(document: Any, text: str, color_hex: str) -> None:
    width = _content_width_cm(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_width(cell, width)
    _shade_cell(cell, color_hex)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _field_line(document: Any, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run_label = paragraph.add_run(f"{label}: ")
    run_label.bold = True
    paragraph.add_run(str(value))


def _field_pair_line(document: Any, label1: str, value1: Any, label2: str, value2: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run1 = paragraph.add_run(f"{label1}: ")
    run1.bold = True
    paragraph.add_run(f"{value1}     ")
    run2 = paragraph.add_run(f"{label2}: ")
    run2.bold = True
    paragraph.add_run(str(value2))


def _fmt_num(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:g}"
    return str(value)


# --------------------------------------------------------------------------------------
# Report sections
# --------------------------------------------------------------------------------------

def _add_header(document: Any, job_name: str) -> None:
    _bar(document, "CODOC RESULTS REPORT", COLOR_TITLE_BAR)
    parts = job_name.split("_")
    date_part = parts[0] if len(parts) > 0 else job_name
    time_part = parts[1] if len(parts) > 1 else ""
    paragraph = document.add_paragraph()
    for label, value in (("Date", date_part), ("Time", time_part), ("Job Name", job_name)):
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(f"{value}     ")


def _add_docking_parameters(document: Any, docking: dict[str, Any], ligand: dict[str, Any]) -> None:
    _bar(document, "DOCKING PARAMETERS:", COLOR_SECTION_BAR)
    _field_line(document, "Score Function", docking.get("scoring_function", ""))
    _field_line(document, "Docking type", docking.get("docking_type", ""))
    _field_line(document, "Processing type", docking.get("processing_type", ""))
    _field_line(document, "Poses number", docking.get("poses", ""))
    _field_line(document, "Minimum RMSD", f"{_fmt_num(docking.get('min_rmsd', ''))} Angstrom")
    _field_line(document, "Energy range", f"{_fmt_num(docking.get('energy_range', ''))} Kcal/mol")
    _field_line(document, "Exhaustiveness", docking.get("exhaustiveness", ""))
    _field_line(document, "Conversion engine", ligand.get("conversion_engine", ""))
    _field_line(document, "Minimization algorithm", ligand.get("minimization_algorithm", ""))
    _field_line(document, "Minimization force field", ligand.get("minimization_forcefield", ""))
    _field_line(document, "Minimization steps", ligand.get("minimization_steps", ""))
    _field_line(document, "Protonation pH", _fmt_num(ligand.get("ph", "")))


def _add_druggability_parameters(document: Any, ligand: dict[str, Any]) -> None:
    _bar(document, "DRUGGABILITY PARAMETERS:", COLOR_SECTION_BAR)
    _field_pair_line(document, "MW min", _fmt_num(ligand.get("mw_min", "")), "MW max", _fmt_num(ligand.get("mw_max", "")))
    _field_pair_line(document, "LogP min", _fmt_num(ligand.get("logp_min", "")), "LogP max", _fmt_num(ligand.get("logp_max", "")))
    _field_pair_line(
        document, "H Donor max", ligand.get("h_donor_max", ""), "H Acceptor max", ligand.get("h_acceptor_max", "")
    )
    _field_pair_line(
        document,
        "Rotatable Bonds max",
        ligand.get("rotatable_bonds_max", ""),
        "TPSA max",
        _fmt_num(ligand.get("tpsa_max", "")),
    )


def _add_conversion_actions(document: Any, completed_actions: set[str]) -> None:
    _bar(document, "CONVERTION LIGANDS ACTIONS:", COLOR_SECTION_BAR)
    for key, label in LIGAND_ACTION_LABELS:
        _field_line(document, label, "Yes" if key in completed_actions else "No")


def _add_databank_summary(document: Any, stats: dict[str, dict[str, int]]) -> None:
    names = list(stats.keys())
    if not names:
        return
    rows = [
        ("Number of initial ligands", "initial"),
        ("Number of rejected ligands:", "rejected"),
        ("Number of empty ligands:", "empty"),
        ("Number of macrocycles fixed:", "macrocycles_fixed"),
        ("Number of recovered or fix ligands", "recovered"),
        ("Number of druggability filtered", "druggability_filtered"),
        ("Total ligands final:", "final"),
    ]
    width = _content_width_cm(document)
    label_width = width * 0.34
    value_width = (width - label_width) / max(len(names), 1)

    table = document.add_table(rows=len(rows) + 1, cols=len(names) + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cell = table.cell(0, 0)
    _set_cell_width(header_cell, label_width)
    _shade_cell(header_cell, COLOR_ACCENT_BAR)
    header_cell.paragraphs[0].add_run("Ligands databank:").bold = True
    for col, name in enumerate(names, start=1):
        cell = table.cell(0, col)
        _set_cell_width(cell, value_width)
        _shade_cell(cell, COLOR_ACCENT_BAR)
        cell.paragraphs[0].add_run(name).bold = True

    for row_idx, (label, key) in enumerate(rows, start=1):
        is_total = key == "final"
        label_cell = table.cell(row_idx, 0)
        _set_cell_width(label_cell, label_width)
        run = label_cell.paragraphs[0].add_run(label)
        run.bold = is_total
        if is_total:
            _shade_cell(label_cell, COLOR_ACCENT_BAR)
        for col, name in enumerate(names, start=1):
            cell = table.cell(row_idx, col)
            _set_cell_width(cell, value_width)
            run = cell.paragraphs[0].add_run(str(stats[name][key]))
            run.bold = is_total
            if is_total:
                _shade_cell(cell, COLOR_ACCENT_BAR)
    document.add_paragraph()


def _add_results_table(document: Any, target: str, databank: str, rows: list[dict[str, str]]) -> None:
    _field_line(document, "Target", target)
    _field_line(document, "Ligands Databank", databank)

    width = _content_width_cm(document)
    col_widths = [width * 0.15, width * 0.30, width * 0.27, width * 0.15, width * 0.13]
    headers = ["LIGAND", "SMILES", "STRUCTURE", "BINDING ENERGY\n(Kcal/mol)", "RMSD (mean)"]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, (header, col_width) in enumerate(zip(headers, col_widths)):
        cell = table.cell(0, col)
        _set_cell_width(cell, col_width)
        for line in header.split("\n"):
            paragraph = cell.paragraphs[0] if cell.paragraphs[0].text == "" else cell.add_paragraph()
            paragraph.add_run(line).bold = True

    for row in rows:
        cells = table.add_row().cells
        _set_cell_width(cells[0], col_widths[0])
        cells[0].paragraphs[0].add_run(row["ligand"])
        _set_cell_width(cells[1], col_widths[1])
        cells[1].paragraphs[0].add_run(row["smiles"])
        _set_cell_width(cells[2], col_widths[2])
        image = molecule_image_bytes(row["smiles"])
        if image is not None:
            cells[2].paragraphs[0].add_run().add_picture(image, width=Cm(col_widths[2] - 0.8))
        else:
            cells[2].paragraphs[0].add_run("N/A")
        _set_cell_width(cells[3], col_widths[3])
        cells[3].paragraphs[0].add_run(row["energy"])
        _set_cell_width(cells[4], col_widths[4])
        cells[4].paragraphs[0].add_run(row["rmsd"])
    document.add_paragraph()


# --------------------------------------------------------------------------------------
# Methodology
# --------------------------------------------------------------------------------------

def _join_bold_list(document_add: Callable[[str, bool], None], items: list[str]) -> None:
    for index, item in enumerate(items):
        if index > 0:
            document_add(", " if index < len(items) - 1 else " and ", False)
        document_add(item, True)


def _add_methodology(
    document: Any,
    docking: dict[str, Any],
    ligand: dict[str, Any],
    databank_names: list[str],
    databank_formats: dict[str, str],
    target_names: list[str],
    total_final_ligands: int,
    top_results: int,
    rmsd_limit: float,
) -> None:
    _bar(document, "DESCRIPTION OF PROPOSED METHODOLOGY", COLOR_ACCENT_BAR)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add(text: str, bold: bool) -> None:
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold

    add("The molecular docking method was performed using ", False)
    add("CODOC", True)
    add(" software, available at ", False)
    add("https://github.com/moimaian/CODOC/tree/main", True)
    add(". For this purpose, the ligand databases ", False)
    _join_bold_list(add, databank_names)
    formats = sorted({fmt for fmt in databank_formats.values() if fmt})
    if formats:
        add(" were initially used in ", False)
        _join_bold_list(add, formats)
        add(" format", False)
    add(" for conversion to .pdbqt through the ", False)
    add(str(ligand.get("conversion_engine", "")), True)
    add(" engine. The ligands' protonation state was adjusted to pH ", False)
    add(_fmt_num(ligand.get("ph", "")), True)
    add(
        ", non-polar hydrogens were removed, partial charges were generated by Gasteiger, and the ",
        False,
    )
    add(str(ligand.get("minimization_algorithm", "")), True)
    add(" algorithm with the ", False)
    add(str(ligand.get("minimization_forcefield", "")), True)
    add(" force field was used over ", False)
    add(str(ligand.get("minimization_steps", "")), True)
    add(
        " steps for 3D geometry minimization and prediction. Druggability filters were applied "
        "and macrocycles were adjusted by removing carbon pseudo-atoms, leaving a total of ",
        False,
    )
    add(str(total_final_ligands), True)
    add(" ligands. The protein targets ", False)
    _join_bold_list(add, target_names)
    add(
        " were prepared by removing water molecules, ions, co-crystallized ligands, and other "
        "heteroatoms that are not standard amino acids. Titratable amino acids had their "
        "protonation state adjusted to pH ",
        False,
    )
    add(_fmt_num(docking.get("target_ph", "")), True)
    add(" through ", False)
    add("PDB2PQR/PROPKA", True)
    add(". The ", False)
    scoring = str(docking.get("scoring_function", ""))
    processing_type = str(docking.get("processing_type", "")).strip().upper()
    algorithm_label = f"{scoring}-GPU" if processing_type == "GPU" else scoring
    add(algorithm_label, True)
    add(" algorithm was used for ", False)
    docking_type_label = str(docking.get("docking_type", "")).strip().lower()
    add(docking_type_label, True)
    add(" docking, with an exhaustiveness of ", False)
    add(str(docking.get("exhaustiveness", "")), True)
    add(", an energy range of ", False)
    add(f"{_fmt_num(docking.get('energy_range', ''))} Kcal/mol", True)
    add(", and a minimum RMSD of ", False)
    add(f"{_fmt_num(docking.get('min_rmsd', ''))} Angstrom", True)
    add(" to generate ", False)
    add(str(docking.get("poses", "")), True)
    add(" poses. The top ", False)
    add(str(top_results), True)
    add(
        " hits were then reported in a table, ranked in decreasing order of binding energy, "
        "with a mean RMSD below ",
        False,
    )
    add(f"{_fmt_num(rmsd_limit)} Angstrom", True)
    add(".", False)


# --------------------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------------------

def generate_final_report(
    job_dir: str,
    job_name: str,
    ligands_dir: str,
    frame: Any,
    docking_settings: dict[str, Any],
    ligand_settings: dict[str, Any],
    top_results: int,
    rmsd_limit: float,
    progress_callback: Optional[Callable[[str], None]] = None,
) -> str:
    """Build <job_dir>/DOCKING/<job_name>_RESULTS_REPORT.docx from the job's rebuilt results
    dataframe (LIGAND, SMILES, LIGAND DATABANK, TARGET, BINDING ENERGY (Kcal/mol), RMSD (mean)),
    settings snapshots, and CONVERSION folder contents. Returns the output path."""
    require_docx()
    require_rdkit()

    def report(text: str) -> None:
        if progress_callback is not None:
            progress_callback(text)

    if frame is None or frame.empty:
        raise RuntimeError("No docking results were found for this job. Run docking and reload the results first.")

    import pandas as pd

    working = frame.copy()
    working["BINDING ENERGY (Kcal/mol)"] = pd.to_numeric(working["BINDING ENERGY (Kcal/mol)"], errors="coerce")
    working["RMSD (mean)"] = pd.to_numeric(working["RMSD (mean)"], errors="coerce")
    working = working.dropna(subset=["BINDING ENERGY (Kcal/mol)", "RMSD (mean)"])
    working = working[working["RMSD (mean)"] < rmsd_limit]

    conversion_dir = os.path.join(job_dir, "CONVERSION")
    docking_dir = os.path.join(job_dir, "DOCKING")
    databank_names = sorted(working["LIGAND DATABANK"].dropna().unique().tolist())
    target_names = sorted(working["TARGET"].dropna().unique().tolist())
    if not databank_names or not target_names:
        raise RuntimeError("No valid docking results (with numeric energy and RMSD) were found for this job.")

    report("Reading ligand preparation folders...")
    databank_stats = compute_databank_stats(ligands_dir, conversion_dir, databank_names)
    databank_formats = {name: detect_databank_format(conversion_dir, ligands_dir, name) for name in databank_names}
    completed_actions = read_completed_actions(conversion_dir)
    total_final_ligands = sum(entry["final"] for entry in databank_stats.values())

    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10)

    _add_header(document, job_name)
    _add_docking_parameters(document, docking_settings, ligand_settings)
    _add_druggability_parameters(document, ligand_settings)
    _add_conversion_actions(document, completed_actions)
    _add_databank_summary(document, databank_stats)

    _bar(document, "DOCKING RESULTS:", COLOR_SECTION_BAR)
    for target in target_names:
        for databank in databank_names:
            subset = working[(working["TARGET"] == target) & (working["LIGAND DATABANK"] == databank)]
            if subset.empty:
                continue
            subset = subset.sort_values(by="BINDING ENERGY (Kcal/mol)", ascending=True).head(top_results)
            report(f"Building table for {target} / {databank} ({len(subset)} ligand(s))...")
            rows = [
                {
                    "ligand": str(record["LIGAND"]),
                    "smiles": str(record["SMILES"]),
                    "energy": f"{record['BINDING ENERGY (Kcal/mol)']:.2f}",
                    "rmsd": f"{record['RMSD (mean)']:.3f}",
                }
                for record in subset.to_dict("records")
            ]
            _add_results_table(document, target, databank, rows)

    _add_methodology(
        document,
        docking_settings,
        ligand_settings,
        databank_names,
        databank_formats,
        target_names,
        total_final_ligands,
        top_results,
        rmsd_limit,
    )

    os.makedirs(docking_dir, exist_ok=True)
    output_path = os.path.join(docking_dir, f"{job_name}_RESULTS_REPORT.docx")
    report(f"Saving {output_path} ...")
    document.save(output_path)
    return output_path


def load_job_settings(job_dir: str) -> tuple[dict[str, Any], dict[str, Any]]:
    """Read the docking_settings.json/ligand_settings.json snapshots written next to a job's
    results, tolerating either of them being absent (older jobs created before this feature)."""
    docking_settings: dict[str, Any] = {}
    ligand_settings: dict[str, Any] = {}
    docking_path = Path(job_dir) / "DOCKING" / "docking_settings.json"
    ligand_path = Path(job_dir) / "CONVERSION" / "ligand_settings.json"
    if docking_path.is_file():
        try:
            docking_settings = json.loads(docking_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            pass
    if ligand_path.is_file():
        try:
            ligand_settings = json.loads(ligand_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            pass
    return docking_settings, ligand_settings
